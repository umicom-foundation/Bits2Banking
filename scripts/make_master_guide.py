# make_master_guide.py
# Generates: C:\Bits2Banking\volumes\Bits_to_Banking_Master_Guide_v1.docx (by default)
# Install dependency: pip install python-docx
# Run (from anywhere):  python C:\Bits2Banking\scripts\make_master_guide.py
# Created by: Sammy Hegab
# Updated: 01-09-2025 (path-safe for C:\Bits2Banking layout)

from pathlib import Path
import argparse
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---- Paths (resolve relative to this script) ----
SCRIPT_PATH = Path(__file__).resolve()
ROOT = SCRIPT_PATH.parents[1]                 # -> C:\Bits2Banking
DEFAULT_OUTDIR = ROOT / "volumes"
DEFAULT_FILENAME = "Bits_to_Banking_Master_Guide_v1.docx"

# ---- Helpers ----
def add_page_numbers(doc: Document):
    sec = doc.sections[-1]
    footer = sec.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE'); p._p.append(fld1)
    p.add_run(" of ")
    fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'), 'NUMPAGES'); p._p.append(fld2)

def ensure_code_style(doc: Document):
    # Monospaced paragraph style + light border used for ASCII diagrams and code
    if "CodeBlockBox" not in [s.name for s in doc.styles]:
        s = doc.styles.add_style("CodeBlockBox", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Consolas"
        s.font.size = Pt(10)

def add_code_box(doc: Document, text: str):
    p = doc.add_paragraph(text, style="CodeBlockBox")
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'8')
        el.set(qn('w:space'),'4')
        el.set(qn('w:color'),'D0D0D0')
        b.append(el)
    pPr.append(b)

def build_master_guide() -> Document:
    doc = Document()
    ensure_code_style(doc)

    # Title + subtitle
    doc.add_paragraph("Bits to Banking").style = doc.styles['Title']
    doc.add_paragraph("Master Guide — From Bits → Banking → Calypso").style = doc.styles['Subtitle']

    # Footer page numbers
    add_page_numbers(doc)

    # Dedication & Acknowledgements
    doc.add_heading("Dedication and Acknowledgements", level=1)
    doc.add_paragraph(
        "First and foremost, all thanks are due to Allah, who blessed us with knowledge, "
        "patience, and strength to pursue this project. I thank my parents, my late father "
        "who always encouraged me to seek knowledge, and my brothers and sisters who "
        "supported me throughout the years."
    )
    doc.add_paragraph(
        "I extend heartfelt thanks to my family, mentors, friends, and colleagues who shaped "
        "my journey of learning and service. To my dear friend Mujahid Sufian and his wife "
        "Dr. Alefiah Mubeen, steadfast supporters of the Umicom Foundation and now contributors "
        "to this book, bringing deep expertise in Machine Learning and Data Science. To "
        "Dr. Ahmed El Banaa, physiotherapy specialist at the Military Hospital in Egypt, who "
        "worked tirelessly with Palestinian patients, and has become a close friend and brother. "
        "To Islam Mahmoud, who despite injury continued designing, documenting and supporting aid "
        "distribution in Gaza. To Mohammed Al Danaf and his family for their bravery and sacrifice "
        "in delivering aid and supporting through their business network in Gaza. To Khaled Matter, "
        "a mobile developer from Gaza, who even while injured volunteered to share knowledge through "
        "his phone. To Mohammed Soliman, always rushing to support financially and contributing his "
        "experience. To Saheesh Rafeeque, business partner and founder of Eastern Bridge, for his "
        "backing and logistics expertise."
    )
    doc.add_paragraph(
        "I also extend heartfelt thanks to Tarneem Elyan, an Artificial Intelligence and Data "
        "Engineer from North Gaza, who despite the hardships of war has volunteered her time to "
        "distribute humanitarian aid in her community and to contribute her knowledge to this book. "
        "Her courage and expertise in data science and AI bring hope and practical value to our shared mission."
    )
    doc.add_paragraph(
        "Finally, to all the brave men and women who risked their lives delivering humanitarian aid, "
        "many martyred or injured, and to the doctors serving Palestinians in Egypt, and to every software "
        "engineer, reviewer, and student who contributes to this vision—this book is dedicated to you."
    )

    # How to Use
    doc.add_heading("How to Use This Series", level=1)
    doc.add_paragraph("• ≤ 4 chapters per volume. Small, printable, and easy to study.")
    doc.add_paragraph("• Kid → Teen → Adult. We begin with everyday analogies, then level up with code and labs.")
    doc.add_paragraph("• Smooth transitions. Each volume ends with a 1-page refresh and “what’s next.”")
    doc.add_paragraph("• Code in boxes; simple diagrams. Monospace, high contrast, print-friendly.")

    # Roadmap
    doc.add_heading("Roadmap at a Glance", level=1)
    roadmap = [
        "Vol. 1 — Foundations of Computing: bits, gates, Boolean, binary & hex, CPU, ISA (x86/ARM/RISC-V), assembly → C → C++ → Java.",
        "Vol. 2 — Operating Systems & Linux From Scratch: what an OS does, kernel vs userland, processes, memory, filesystems, network; build a tiny Linux.",
        "Vol. 3 — Databases & Data Engineering: PostgreSQL/Oracle, indexing/transactions, NoSQL, warehousing, ETL, data quality.",
        "Vol. 4 — Networking & Protocols: TCP/UDP, HTTP/REST, TLS; FIX, SWIFT ISO 20022, FpML; queues and messaging.",
        "Vol. 5 — Cybersecurity for Builders: threats, OWASP, memory bugs, secrets/KMS, secure deployment, audit/compliance.",
        "Vol. 6 — Programming Deep Dives: C/C++, Java (JVM, concurrency), Python, Rust, Zig, PHP; patterns, testing, packaging, web & mobile overviews.",
        "Vol. 7 — Finance & Treasury Primer: markets 101, asset classes, front/middle/back office, P&L, risk, settlement, clearing; history of money → digital.",
        "Vols. 8–12 — Calypso Series (multiple slim books): Fast-Track; Architecture & Core; Risk Compute; Position/Liquidity/KPI; Market Data/Quotes; Integration/Messaging.",
        "Vol. 13 — Customization & Extensions: SPIs, adapters, scheduled tasks, UI add-ons, reports; multi-language examples.",
        "Vol. 14 — Testing, Deployment & Ops: CATT, perf/load, observability, HA/DR, upgrades, blue/green, rollback.",
        "Vol. 15 — RISC-V & Advanced Computing: assembly → Linux on RISC-V, QEMU/WSL2 labs, cross-compiling, tiny compiler projects.",
        "Vol. 16 — Projects & Capstones: commodity exchange, privacy-first video platform, social + fundraising site, UmiCOIN (asset-backed)."
    ]
    for line in roadmap:
        doc.add_paragraph("• " + line)

    # Mini-Introductions (samples)
    doc.add_heading("Mini-Introductions (Sample)", level=1)

    doc.add_heading("Volume 1 — Foundations of Computing", level=2)
    doc.add_paragraph(
        "What you’ll get: the “why” behind every 1 and 0. Play with binary using fingers, build AND/OR with paper "
        "switches, watch a CPU “fetch-decode-execute,” then translate simple programs from Python → C → assembly."
    )
    doc.add_paragraph("You’ll be able to:")
    doc.add_paragraph("1) Explain binary and overflow to a kid.")
    doc.add_paragraph("2) Write a 10-line program and peek at its assembly.")
    doc.add_paragraph("3) Connect CPU cycles to how a trade instruction is processed.")

    doc.add_heading("Volume 2 — Operating Systems & Linux From Scratch", level=2)
    doc.add_paragraph(
        "We explain processes, memory, files, users, then compile a tiny Linux. You’ll boot your own system and run basic tools."
    )
    doc.add_paragraph("You’ll be able to:")
    doc.add_paragraph("1) Read process lists and memory usage.")
    doc.add_paragraph("2) Understand system services (systemd) and logs.")
    doc.add_paragraph("3) Prepare a finance-friendly Linux profile.")

    # Video micro-lessons
    doc.add_heading("Video Micro-Lessons (Titles & Descriptions)", level=1)
    videos = [
        ("Bits vs Switches (3 min)", "A light switch becomes a 1 or 0; two switches make AND/OR."),
        ("Binary Counting with Fingers (5 min)", "Count to 31 on one hand; meet overflow."),
        ("From CPU to Instructions (6 min)", "Fetch, decode, execute; registers are sticky notes."),
        ("Hello, Assembly (7 min)", "Add two numbers with registers & memory moves."),
        ("C to Machine Code (8 min)", "Compile a tiny C program and compare assembly."),
        ("Processes & Memory (6 min)", "Stacks/heaps/syscalls; why programs crash."),
        ("Build a Tiny Linux (8 min)", "Compile + boot a minimal system; first prompt!"),
        ("Your First REST Endpoint (5 min)", "Create /hello; add a nightly job.")
    ]
    for t, d in videos:
        doc.add_paragraph(f"• {t} — {d}")

    # Lightweight ASCII diagrams in bordered boxes
    doc.add_heading("Lightweight Diagrams (Printable ASCII)", level=1)
    add_code_box(doc, """
[CPU]--fetch-->[Instruction]--decode-->[Execute]
   |                               |
 [Registers] <-------------------- [ALU]
""".strip("\n"))
    add_code_box(doc, """
Process
  ├─ Code (read-only)
  ├─ Heap (grows up)     <-- new objects
  └─ Stack (grows down)  <-- function calls
""".strip("\n"))
    add_code_box(doc, r"""
REST Control Plane
Client -> Risk Services -> Risk Server
   \__ health/status      \__ MQ + compute
""".strip("\n"))

    # On-ramp labs
    doc.add_heading("On-Ramp Labs (Quick Wins)", level=1)
    doc.add_paragraph("• Build & run a one-file program in C, C++, Java, Python.")
    doc.add_paragraph("• Boot a tiny Linux VM; list processes & services.")
    doc.add_paragraph("• Create /hello in Java; schedule a nightly job; check logs & metrics.")
    doc.add_paragraph("• Create a DB table, insert a row, read it back.")

    # Glossary
    doc.add_heading("Glossary (Short & Sweet)", level=1)
    glossary = [
        ("Bit", "Tiniest piece of info: 0 or 1."),
        ("CPU", "The ‘thinker’ that follows instructions fast."),
        ("Register", "A tiny ultra-fast box inside the CPU."),
        ("Process", "A running program with its own memory world."),
        ("Message Queue", "A mailbox for servers to pass work."),
        ("Latency", "How long something takes (track p95/p99).")
    ]
    for k, v in glossary:
        doc.add_paragraph(f"• {k} — {v}")

    # Next steps
    doc.add_heading("Where to Go Next", level=1)
    doc.add_paragraph(
        "Start with Volume 1 — Foundations of Computing, then Volume 2 — Operating Systems & Linux From Scratch. "
        "When ready, enter the Calypso volumes for banking-grade platforms."
    )

    return doc

def main():
    parser = argparse.ArgumentParser(description="Build the Bits to Banking Master Guide DOCX.")
    parser.add_argument("--outdir", default=str(DEFAULT_OUTDIR), help="Output directory (default: volumes folder)")
    parser.add_argument("--filename", default=DEFAULT_FILENAME, help="Output filename")
    args = parser.parse_args()

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    out_path = outdir / args.filename

    doc = build_master_guide()
    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == "__main__":
    main()

