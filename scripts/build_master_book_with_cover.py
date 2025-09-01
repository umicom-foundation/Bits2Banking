# build_master_book_with_cover.py
# Merge built volumes into book\Bits_to_Banking_Master_Book.docx
# Adds full-page cover (no page number on cover). Fixes: no leading blank page, no duplicate page numbers.
# Created by: Sammy Hegab  Date: 2025-09-01

from pathlib import Path
from datetime import datetime
import argparse, json
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"C:\Bits2Banking")
BOOK_DIR = ROOT / "book"
VOLUMES_DIR = ROOT / "volumes"
TOC_PATH = ROOT / "toc.json"

def find_cover(user_path: str | None):
    if user_path:
        p = Path(user_path)
        if p.exists():
            return p
    # Prefer generated A4 cover
    p = ROOT / "bookcover" / "Bits_to_Banking_Cover_A4_300dpi.png"
    if p.exists():
        return p
    # Fall back to common names in images/
    for name in ("cover.png", "cover.jpg", "cover.jpeg", "Book cover for bit to Banking.png"):
        p = ROOT / "images" / name
        if p.exists():
            return p
    return None

def _remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)

def clear_footer(section):
    footer = section.footer
    for p in list(footer.paragraphs):
        _remove_paragraph(p)

def add_footer_page_numbers(section, first_page_blank: bool):
    # Make this section's footer independent and tidy
    section.footer.is_linked_to_previous = False
    section.different_first_page_header_footer = first_page_blank
    clear_footer(section)
    p = section.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    f1 = OxmlElement("w:fldSimple"); f1.set(qn("w:instr"), "PAGE"); p._p.append(f1)
    p.add_run(" of ")
    f2 = OxmlElement("w:fldSimple"); f2.set(qn("w:instr"), "NUMPAGES"); p._p.append(f2)

def restart_page_numbering_at(section, start_at: int = 1):
    sectPr = section._sectPr
    for child in list(sectPr):
        if child.tag == qn("w:pgNumType"):
            sectPr.remove(child)
    pg = OxmlElement("w:pgNumType")
    pg.set(qn("w:start"), str(start_at))
    sectPr.append(pg)

def append_body(dst: Document, src: Document):
    # Append the body elements from src into dst
    for block in src.element.body:
        dst.element.body.append(block)

def load_volumes_list():
    vols = []
    if TOC_PATH.exists():
        toc = json.loads(TOC_PATH.read_text(encoding="utf-8-sig"))  # tolerate BOM
        for v in toc.get("volumes", []):
            vols.append(VOLUMES_DIR / v["output"])
    else:
        # Optional fallback: stitch all DOCX in volumes/ if no TOC
        # vols = sorted(VOLUMES_DIR.glob("*.docx"))
        pass
    return vols

def safe_save(doc, out_path: Path):
    out_path.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(out_path)
        print("Created:", out_path)
    except PermissionError:
        alt = out_path.with_name(f"{out_path.stem}_{datetime.now().strftime('%Y%m%d_%H%M%S')}{out_path.suffix}")
        doc.save(alt)
        print(f"Target in use; saved as: {alt}")

def build_master(cover_path: Path, out_path: Path):
    BOOK_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()

    # Remove default empty first paragraph so the cover is truly first content
    if len(doc.paragraphs) == 1 and not doc.paragraphs[0].text:
        _remove_paragraph(doc.paragraphs[0])

    # COVER section — zero margins, no page number printed
    cover_sec = doc.sections[0]
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(cover_sec, side, Inches(0))

    add_footer_page_numbers(cover_sec, first_page_blank=True)

    # Insert the cover image (full width) and center the paragraph
    if cover_path and cover_path.exists():
        doc.add_picture(str(cover_path), width=cover_sec.page_width)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        title = doc.add_paragraph("Bits to Banking"); title.style = doc.styles["Title"]
        sub = doc.add_paragraph("Master Book"); sub.style = doc.styles["Subtitle"]

    # CONTENT section — normal margins, restart numbering at 1, clean footer
    content_sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
    content_sec.top_margin = content_sec.bottom_margin = Inches(1)
    content_sec.left_margin = content_sec.right_margin = Inches(1)
    content_sec.footer.is_linked_to_previous = False
    restart_page_numbering_at(content_sec, 1)
    add_footer_page_numbers(content_sec, first_page_blank=False)

    # Append volumes in order — use the content section for the first one
    vols = load_volumes_list()
    if not vols:
        raise SystemExit(
            f"No volumes found.\n"
            f" - Build volumes into {VOLUMES_DIR}\n"
            f" - Or create {TOC_PATH} listing them (key 'output')."
        )

    for idx, vol_docx in enumerate(vols):
        if not vol_docx.exists():
            raise SystemExit(f"Missing built volume:\n  {vol_docx}")

        if idx > 0:
            # Subsequent volumes start on a new page, inherit footer to avoid duplicates
            sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
            sec.footer.is_linked_to_previous = True

        append_body(doc, Document(str(vol_docx)))

    safe_save(doc, out_path)

def main():
    ap = argparse.ArgumentParser(description="Build master book with full-page cover and clean footers")
    ap.add_argument("--cover", help="Custom cover path (optional)")
    ap.add_argument("--out", default=str(BOOK_DIR / "Bits_to_Banking_Master_Book.docx"))
    args = ap.parse_args()

    cover = find_cover(args.cover)
    build_master(cover, Path(args.out))

if __name__ == "__main__":
    main()
