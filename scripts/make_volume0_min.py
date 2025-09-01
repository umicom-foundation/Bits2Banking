# make_volume0_min.py
# Purpose: Create a tiny Volume 0 Word document (no covers, no extras) so we have a clean first success.
# Created by: Sammy Hegab
# Date: 2025-09-01

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

OUT_PATH = Path(r"C:\Bits2Banking\volumes\Volume_00_Source_Control.docx")

def add_page_numbers(doc: Document):
    sec = doc.sections[-1]
    # one centered "Page X of Y" in the footer
    p = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    f1 = OxmlElement("w:fldSimple"); f1.set(qn("w:instr"), "PAGE"); p._p.append(f1)
    p.add_run(" of ")
    f2 = OxmlElement("w:fldSimple"); f2.set(qn("w:instr"), "NUMPAGES"); p._p.append(f2)

def main():
    doc = Document()

    # Title page
    title = doc.add_paragraph("Volume 0 — Source Control & Collaboration")
    title.style = doc.styles["Title"]
    subtitle = doc.add_paragraph("Simple starter build (no cover, no extras)")
    subtitle.style = doc.styles["Subtitle"]

    add_page_numbers(doc)

    # Chapter 0: Prelude (short text — we can expand later)
    doc.add_page_break()
    doc.add_heading("Prelude — Recap & Setup", level=1)
    doc.add_paragraph(
        "In this volume we set up Git, GitHub, Pull Requests, reviews, and a tiny workflow "
        "so we can build our book step by step."
    )
    doc.add_paragraph("Prerequisites: Windows 10/11, internet, Python 3, Git, and PowerShell.")

    # Chapter 1: Git & GitHub basics (short text — we’ll grow later)
    doc.add_page_break()
    doc.add_heading("Chapter 1 — Source Control 101: Git & GitHub (Web + CLI)", level=1)
    doc.add_paragraph("Goal: clone a repo, make a branch, commit a change, and open a Pull Request.")

    # Wrap-up
    doc.add_page_break()
    doc.add_heading("Wrap-Up — What You Can Do & What’s Next", level=1)
    doc.add_paragraph(
        "You created your first volume file! Next we’ll add more chapters and automation, "
        "but this proves the basic build works."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    print("Created:", OUT_PATH)

if __name__ == "__main__":
    main()

