# make_volume0_from_md.py
# Purpose: Build Volume 0 DOCX from two markdown files (simple headings, lists, code fences)
# Created by: Sammy Hegab
# Date: 2025-09-01

from pathlib import Path
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# --------------------------
# IMPORTANT CHANGE: paths
# --------------------------
# Use repo-relative paths so it works on GitHub runners and your PC.
ROOT = Path(__file__).resolve().parents[1]  # repo root = parent of /scripts
CH0 = ROOT / "chapters" / "v00" / "ch00_prelude.md"
CH1 = ROOT / "chapters" / "v00" / "ch01_git_GitHub_basics.md"
OUT = ROOT / "volumes" / "Volume_00_Source_Control.docx"

IMG_RE = re.compile(r'!\[(.*?)\]\((.*?)\)')  # (images ignored for now)

def ensure_code_style(doc: Document):
    names = [s.name for s in doc.styles]
    if "CodeBlockBox" not in names:
        s = doc.styles.add_style("CodeBlockBox", WD_STYLE_TYPE.PARAGRAPH)
        s.font.name = "Consolas"
        s.font.size = Pt(10)

def add_code_box(doc: Document, text: str):
    p = doc.add_paragraph(text, style="CodeBlockBox")
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:pBdr')
    for side in ('top','left','bottom','right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'),'single')
        el.set(qn('w:sz'),'8')
        el.set(qn('w:space'),'4')
        el.set(qn('w:color'),'D0D0D0')
        b.append(el)
    pPr.append(b)

def add_page_numbers(doc: Document):
    sec = doc.sections[-1]
    p = sec.footer.paragraphs[0] if sec.footer.paragraphs else sec.footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Page ")
    f1 = OxmlElement("w:fldSimple"); f1.set(qn("w:instr"), "PAGE"); p._p.append(f1)
    p.add_run(" of ")
    f2 = OxmlElement("w:fldSimple"); f2.set(qn("w:instr"), "NUMPAGES"); p._p.append(f2)

def add_md_chapter(doc: Document, md_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    # use first line that starts with '# ' as the chapter heading
    heading = None
    body_start = 0
    for i, line in enumerate(lines):
        if line.startswith("# "):
            heading = line[2:].strip()
            body_start = i + 1
            break
    if not heading:
        heading = md_path.stem.replace("_", " ").title()

    doc.add_page_break()
    doc.add_heading(heading, level=1)

    # simple markdown: #/##/### headings, bullets, numbers, code fences
    in_code, buf = False, []
    def flush_code():
        nonlocal buf
        if buf:
            add_code_box(doc, "\n".join(buf))
            buf = []

    for line in lines[body_start:]:
        if line.strip().startswith("```"):
            if not in_code:
                in_code, buf = True, []
            else:
                in_code = False
                flush_code()
            continue
        if in_code:
            buf.append(line)
            continue

        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif re.match(r'^\s*[-*]\s+', line):
            doc.add_paragraph(line.strip()[2:], style='List Bullet')
        elif re.match(r'^\s*\d+\.\s+', line):
            doc.add_paragraph(line.strip(), style='List Number')
        elif not line.strip():
            doc.add_paragraph("")
        else:
            # ignore images for now to keep it simple
            if IMG_RE.match(line.strip()):
                doc.add_paragraph("[image omitted in this simple build]")
            else:
                doc.add_paragraph(line)

def main():
    # make sure chapter files exist
    for p in (CH0, CH1):
        if not p.exists():
            raise SystemExit(f"Missing: {p}")

    doc = Document()
    ensure_code_style(doc)

    # simple title page
    t = doc.add_paragraph("Volume 0 — Source Control & Collaboration")
    t.style = doc.styles["Title"]
    s = doc.add_paragraph("Simple build from two markdown chapters")
    s.style = doc.styles["Subtitle"]

    add_page_numbers(doc)

    # add both chapters
    add_md_chapter(doc, CH0)
    add_md_chapter(doc, CH1)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print("Built:", OUT)

if __name__ == "__main__":
    main()
# End of make_volume0_from_md.py